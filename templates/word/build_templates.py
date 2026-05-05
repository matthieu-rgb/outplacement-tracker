"""
build_templates.py
------------------
Generates transfer_mappe_template_de.docx and transfer_mappe_template_en.docx
with the 10 k Beratung visual identity.

Each file contains 118 Plain Text Content Controls (w:sdt elements) with the
exact Tag values required by the Power Automate "Populate a Microsoft Word
template" action (Word Online Business connector).

Visual identity:
  - BLUE_STEEL  (#3D81B0) : main titles, field underlines, header accent
  - BLUE_DEEP   (#2A5A7A) : sub-headings
  - GREY_DARK   (#333333) : body text
  - GREY_LIGHT  (#9CA3AF) : left column labels, header/footer text

Typography (fallback chain applied at render time by Word / LibreOffice):
  - Main titles       : Georgia 28pt bold BLUE_STEEL
  - Section headings  : Georgia 16pt bold BLUE_STEEL
  - Sub-headings      : Arial 12pt BLUE_DEEP
  - Labels left col   : Arial 10pt GREY_LIGHT
  - Body / SDT text   : Georgia 11pt GREY_DARK

Layout:
  - A4, margins top/bottom/left 2.5cm, right 2.0cm, usable width 16.5cm
  - Invisible 2-column table: left 4.0cm (labels), right 12.5cm (SDTs)
  - Field bottom border: single 0.5pt #3D81B0

Header (all pages): chapter title left | "10 k" right | bottom rule #9CA3AF
Footer (all pages): branding left | "Seite/Page [PAGE]" right

Dependencies:
    pip install python-docx

Usage:
    python3 build_templates.py

Output:
    templates/word/transfer_mappe_template_de.docx
    templates/word/transfer_mappe_template_en.docx

Verification:
    python3 -c "
    from docx import Document
    import zipfile, re
    d = Document('templates/word/transfer_mappe_template_de.docx')
    with zipfile.ZipFile('templates/word/transfer_mappe_template_de.docx') as z:
        xml = z.read('word/document.xml').decode()
    tags = re.findall(r'<w:tag w:val=\"([^\"]+)\"', xml)
    print(len(tags), 'tag values found')
    "
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Palette 10 k Beratung
# ---------------------------------------------------------------------------

BLUE_STEEL = RGBColor(0x3D, 0x81, 0xB0)
BLUE_DEEP  = RGBColor(0x2A, 0x5A, 0x7A)
GREY_DARK  = RGBColor(0x33, 0x33, 0x33)
GREY_LIGHT = RGBColor(0x9C, 0xA3, 0xAF)

# Hex strings for OxmlElement color attributes (no leading #)
HEX_BLUE_STEEL = "3D81B0"
HEX_GREY_LIGHT = "9CA3AF"

# Column widths in cm
LEFT_COL_CM  = 4.0
RIGHT_COL_CM = 12.5


# ---------------------------------------------------------------------------
# Content Control tag lists (FROZEN - must not change)
# ---------------------------------------------------------------------------

COVER_TAGS = [
    "doc_titre",
    "participant_prenom",
    "participant_nom",
    "participant_date_debut",
    "conseillere_nom",
    "doc_date_generation",
]

PROFIL_TAGS = [
    "profil_plan_a",
    "profil_plan_b",
    "profil_marketingplan",
    "profil_zielmarkt",
]

BILAN_FIELDS = [
    "date_rdv",
    "date_soumission",
    "bilan_general",
    "statut_objectifs",
    "statut_objectifs_detail",
    "was_lief_gut",
    "wo_brauche_ich",
    "themen_naechster_termin",
    "sonstige_anmerkungen",
]

# Expected total: 6 + 4 + (12 * 9) = 118


# ---------------------------------------------------------------------------
# DE template content
# ---------------------------------------------------------------------------

DE_STRINGS = {
    "doc_title_fixed":  "TRANSFER MAPPE",
    "doc_titre_placeholder": "Transfer Mappe",
    "confidential": "Vertraulich - Nur fuer den internen Gebrauch",
    "cover_labels": {
        "doc_titre":              "Dokumenttitel",
        "participant_prenom":     "Teilnehmer/in - Vorname",
        "participant_nom":        "Teilnehmer/in - Nachname",
        "participant_date_debut": "Beginn des Parcours",
        "conseillere_nom":        "Beraterin",
        "doc_date_generation":    "Erstellt am",
    },
    "section_profil_heading":  "1. Karriereprofil",
    "section_profil_sub":      "Berufliche Zielsetzung",
    "profil_labels": {
        "profil_plan_a":         "Plan A - Berufliches Hauptziel",
        "profil_plan_b":         "Plan B - Berufliches Alternativziel",
        "profil_marketingplan":  "Berufliches Profil und Staerken (Marketingplan)",
        "profil_zielmarkt":      "Zielmarkt",
    },
    "bilan_heading":  "Monatsbericht {nn:02d}",
    "bilan_sub":      "Zielvereinbarung",
    "bilan_labels": {
        "date_rdv":                "Datum des Termins",
        "date_soumission":         "Eingangsdatum",
        "bilan_general":           "Allgemeine Einschaetzung des Monats",
        "statut_objectifs":        "Status der vorherigen Ziele",
        "statut_objectifs_detail": "Details zum Zielstatus",
        "was_lief_gut":            "Was lief gut?",
        "wo_brauche_ich":          "Wo brauche ich Unterstuetzung?",
        "themen_naechster_termin": "Themen fuer den naechsten Termin",
        "sonstige_anmerkungen":    "Sonstige Anmerkungen",
    },
    "signature": {
        "title":       "Zielvereinbarung - Unterschriften",
        "datum":       "Datum: .............................",
        "left_label":  "Teilnehmer/in:",
        "right_label": "                              Beraterin:",
        "left_line":   "_________________________________",
        "right_line":  "       _________________________________",
    },
    "header_left":   "Transfer Mappe",
    "footer_left":   "10 k Beratung GmbH  |  Transfer Mappe",
    "lang":          "DE",
}


# ---------------------------------------------------------------------------
# EN template content
# ---------------------------------------------------------------------------

EN_STRINGS = {
    "doc_title_fixed":  "TRANSFER PORTFOLIO",
    "doc_titre_placeholder": "Transfer Portfolio",
    "confidential": "Confidential - For internal use only",
    "cover_labels": {
        "doc_titre":              "Document title",
        "participant_prenom":     "Participant - First name",
        "participant_nom":        "Participant - Last name",
        "participant_date_debut": "Start date",
        "conseillere_nom":        "Advisor",
        "doc_date_generation":    "Generated on",
    },
    "section_profil_heading":  "1. Career Profile",
    "section_profil_sub":      "Professional objectives",
    "profil_labels": {
        "profil_plan_a":         "Plan A - Primary professional objective",
        "profil_plan_b":         "Plan B - Alternative professional objective",
        "profil_marketingplan":  "Professional profile and strengths (Marketing plan)",
        "profil_zielmarkt":      "Target market",
    },
    "bilan_heading":  "Monthly Update {nn:02d}",
    "bilan_sub":      "Monthly Review",
    "bilan_labels": {
        "date_rdv":                "Appointment date",
        "date_soumission":         "Submission date",
        "bilan_general":           "General monthly review",
        "statut_objectifs":        "Previous objectives status",
        "statut_objectifs_detail": "Objectives status detail",
        "was_lief_gut":            "What went well?",
        "wo_brauche_ich":          "Where do I need support?",
        "themen_naechster_termin": "Topics for next appointment",
        "sonstige_anmerkungen":    "Additional remarks",
    },
    "signature": {
        "title":       "Monthly Review - Signatures",
        "datum":       "Date: .............................",
        "left_label":  "Participant:",
        "right_label": "                              Advisor:",
        "left_line":   "_________________________________",
        "right_line":  "       _________________________________",
    },
    "header_left":   "Transfer Portfolio",
    "footer_left":   "10 k Beratung GmbH  |  Transfer Portfolio",
    "lang":          "EN",
}


# ---------------------------------------------------------------------------
# SDT builder - bottom border (underline bleu acier)
# ---------------------------------------------------------------------------

def make_sdt_with_underline(tag_val, placeholder=""):
    """
    Build a Plain Text Content Control (w:sdt) with a blue steel bottom border.
    The underline bottom border is applied to the inner paragraph via w:pBdr/w:bottom.
    Tag value is the exact key used by Power Automate for field matching.
    """
    sdt = OxmlElement('w:sdt')

    # --- sdtPr: tag, alias, text (Plain Text) ---
    sdtPr = OxmlElement('w:sdtPr')

    tag = OxmlElement('w:tag')
    tag.set(qn('w:val'), tag_val)
    sdtPr.append(tag)

    alias = OxmlElement('w:alias')
    alias.set(qn('w:val'), tag_val)
    sdtPr.append(alias)

    text_el = OxmlElement('w:text')
    sdtPr.append(text_el)

    sdt.append(sdtPr)

    # --- sdtContent: inner paragraph with bottom border + run ---
    sdtContent = OxmlElement('w:sdtContent')
    inner_p = OxmlElement('w:p')

    # Paragraph properties: bottom border blue steel 0.5pt
    pPr = OxmlElement('w:pPr')
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')      # 4 eighths-of-a-point = 0.5pt
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), HEX_BLUE_STEEL)
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Run properties: Georgia 11pt GREY_DARK
    rPr_inner = OxmlElement('w:rPr')
    rFonts_inner = OxmlElement('w:rFonts')
    rFonts_inner.set(qn('w:ascii'), 'Georgia')
    rFonts_inner.set(qn('w:hAnsi'), 'Georgia')
    rPr_inner.append(rFonts_inner)
    sz_inner = OxmlElement('w:sz')
    sz_inner.set(qn('w:val'), '22')   # 22 half-points = 11pt
    rPr_inner.append(sz_inner)
    color_inner = OxmlElement('w:color')
    color_inner.set(qn('w:val'), '333333')
    rPr_inner.append(color_inner)

    inner_p.append(pPr)

    inner_r = OxmlElement('w:r')
    inner_r.append(rPr_inner)
    inner_t = OxmlElement('w:t')
    inner_t.text = placeholder
    inner_r.append(inner_t)
    inner_p.append(inner_r)

    sdtContent.append(inner_p)
    sdt.append(sdtContent)

    return sdt


# ---------------------------------------------------------------------------
# Invisible 2-column table helpers
# ---------------------------------------------------------------------------

def create_two_col_table(doc, left_cm=LEFT_COL_CM, right_cm=RIGHT_COL_CM):
    """
    Create an invisible 2-column table for the label/value layout.
    All borders are suppressed. Column widths are set in DXA (twips).
    """
    table = doc.add_table(rows=0, cols=2)
    tbl = table._tbl

    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Total width - remove any existing tblW before appending
    existing_tblW = tblPr.find(qn('w:tblW'))
    if existing_tblW is not None:
        tblPr.remove(existing_tblW)
    tblW = OxmlElement('w:tblW')
    total_twips = int((left_cm + right_cm) * 567)
    tblW.set(qn('w:w'), str(total_twips))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

    # Remove all borders - replace existing tblBorders if present
    existing_tblBorders = tblPr.find(qn('w:tblBorders'))
    if existing_tblBorders is not None:
        tblPr.remove(existing_tblBorders)
    tblBorders = OxmlElement('w:tblBorders')
    for bname in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{bname}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        tblBorders.append(b)
    tblPr.append(tblBorders)

    # Column grid - remove existing tblGrid first (python-docx may have created one)
    existing_tblGrid = tbl.find(qn('w:tblGrid'))
    if existing_tblGrid is not None:
        tbl.remove(existing_tblGrid)
    tblGrid = OxmlElement('w:tblGrid')
    for w_cm in [left_cm, right_cm]:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(int(w_cm * 567)))
        tblGrid.append(gc)

    tblPr_idx = list(tbl).index(tblPr)
    tbl.insert(tblPr_idx + 1, tblGrid)

    return table


def set_cell_width(cell, w_cm):
    """Set cell width in cm (converted to DXA twips)."""
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:tcW'))
    if existing is not None:
        tcPr.remove(existing)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(w_cm * 567)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


def add_field_row(table, label, tag_val):
    """
    Add a table row: left cell = grey label (Arial 10pt), right cell = SDT with underline.
    """
    row = table.add_row()
    left_cell, right_cell = row.cells[0], row.cells[1]

    set_cell_width(left_cell, LEFT_COL_CM)
    set_cell_width(right_cell, RIGHT_COL_CM)

    # Left cell: label
    p_left = left_cell.paragraphs[0]
    run_label = p_left.add_run(label)
    run_label.font.size = Pt(10)
    run_label.font.color.rgb = GREY_LIGHT
    run_label.font.name = 'Arial'

    # Right cell: SDT only (clear default paragraph, insert raw SDT element)
    for child in list(right_cell._tc):
        right_cell._tc.remove(child)
    right_cell._tc.append(make_sdt_with_underline(tag_val))

    return row


# ---------------------------------------------------------------------------
# Header and footer
# ---------------------------------------------------------------------------

def set_header(section, left_text, right_text="10 k"):
    """
    Set page header: chapter title left, '10 k' right, grey bottom rule.
    Uses a tab stop at the full usable width for right-alignment.
    """
    header = section.header

    # Clear any existing content
    for p in header.paragraphs:
        p._p.getparent().remove(p._p)

    hdr_elem = header._element   # <w:hdr> element
    p_elem = OxmlElement('w:p')
    hdr_elem.append(p_elem)

    # Paragraph properties: right tab stop + bottom border
    pPr = OxmlElement('w:pPr')

    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), str(int(16.5 * 567)))  # 16.5cm in twips
    tabs.append(tab)
    pPr.append(tabs)

    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), HEX_GREY_LIGHT)
    pBdr.append(bottom)
    pPr.append(pBdr)

    p_elem.append(pPr)

    # Left run: chapter title, Arial 9pt GREY_LIGHT
    r_left = OxmlElement('w:r')
    rPr_left = OxmlElement('w:rPr')
    rFonts_left = OxmlElement('w:rFonts')
    rFonts_left.set(qn('w:ascii'), 'Arial')
    rFonts_left.set(qn('w:hAnsi'), 'Arial')
    rPr_left.append(rFonts_left)
    sz_left = OxmlElement('w:sz')
    sz_left.set(qn('w:val'), '18')   # 18 half-points = 9pt
    rPr_left.append(sz_left)
    color_left = OxmlElement('w:color')
    color_left.set(qn('w:val'), HEX_GREY_LIGHT)
    rPr_left.append(color_left)
    r_left.append(rPr_left)
    t_left = OxmlElement('w:t')
    t_left.text = left_text
    r_left.append(t_left)
    p_elem.append(r_left)

    # Tab run
    r_tab = OxmlElement('w:r')
    t_tab = OxmlElement('w:tab')
    r_tab.append(t_tab)
    p_elem.append(r_tab)

    # Right run: "10 k", Georgia 11pt BLUE_STEEL bold
    r_right = OxmlElement('w:r')
    rPr_right = OxmlElement('w:rPr')
    b_right = OxmlElement('w:b')
    rPr_right.append(b_right)
    rFonts_right = OxmlElement('w:rFonts')
    rFonts_right.set(qn('w:ascii'), 'Georgia')
    rFonts_right.set(qn('w:hAnsi'), 'Georgia')
    rPr_right.append(rFonts_right)
    sz_right = OxmlElement('w:sz')
    sz_right.set(qn('w:val'), '22')  # 22 half-points = 11pt
    rPr_right.append(sz_right)
    color_right = OxmlElement('w:color')
    color_right.set(qn('w:val'), HEX_BLUE_STEEL)
    rPr_right.append(color_right)
    r_right.append(rPr_right)
    t_right = OxmlElement('w:t')
    t_right.text = right_text
    r_right.append(t_right)
    p_elem.append(r_right)


def set_footer(section, left_text, lang="DE"):
    """
    Set page footer: branding text left, page number right.
    "Seite [PAGE]" (DE) or "Page [PAGE]" (EN).
    """
    footer = section.footer

    for p in footer.paragraphs:
        p._p.getparent().remove(p._p)

    ftr_elem = footer._element   # <w:ftr> element
    p_elem = OxmlElement('w:p')
    ftr_elem.append(p_elem)

    # Paragraph properties: right tab stop + top border separator
    pPr = OxmlElement('w:pPr')
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), str(int(16.5 * 567)))
    tabs.append(tab)
    pPr.append(tabs)
    pBdr_ftr = OxmlElement('w:pBdr')
    top_ftr = OxmlElement('w:top')
    top_ftr.set(qn('w:val'), 'single')
    top_ftr.set(qn('w:sz'), '4')
    top_ftr.set(qn('w:space'), '1')
    top_ftr.set(qn('w:color'), 'E5E5E5')
    pBdr_ftr.append(top_ftr)
    pPr.append(pBdr_ftr)
    p_elem.append(pPr)

    # Left run: branding, Arial 9pt GREY_LIGHT
    r_left = OxmlElement('w:r')
    rPr_left = OxmlElement('w:rPr')
    rFonts_left = OxmlElement('w:rFonts')
    rFonts_left.set(qn('w:ascii'), 'Arial')
    rFonts_left.set(qn('w:hAnsi'), 'Arial')
    rPr_left.append(rFonts_left)
    sz_left = OxmlElement('w:sz')
    sz_left.set(qn('w:val'), '18')
    rPr_left.append(sz_left)
    color_left = OxmlElement('w:color')
    color_left.set(qn('w:val'), HEX_GREY_LIGHT)
    rPr_left.append(color_left)
    r_left.append(rPr_left)
    t_left = OxmlElement('w:t')
    t_left.text = left_text
    r_left.append(t_left)
    p_elem.append(r_left)

    # Tab
    r_tab = OxmlElement('w:r')
    t_tab_el = OxmlElement('w:tab')
    r_tab.append(t_tab_el)
    p_elem.append(r_tab)

    # Page label run: Arial 9pt GREY_LIGHT
    page_label = "Seite " if lang == "DE" else "Page "
    r_label = OxmlElement('w:r')
    rPr_label = OxmlElement('w:rPr')
    rFonts_label = OxmlElement('w:rFonts')
    rFonts_label.set(qn('w:ascii'), 'Arial')
    rFonts_label.set(qn('w:hAnsi'), 'Arial')
    rPr_label.append(rFonts_label)
    sz_label = OxmlElement('w:sz')
    sz_label.set(qn('w:val'), '18')
    rPr_label.append(sz_label)
    color_label = OxmlElement('w:color')
    color_label.set(qn('w:val'), HEX_GREY_LIGHT)
    rPr_label.append(color_label)
    r_label.append(rPr_label)
    t_label = OxmlElement('w:t')
    t_label.text = page_label
    t_label.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r_label.append(t_label)
    p_elem.append(r_label)

    # PAGE field: begin instrText end
    rPr_fld = OxmlElement('w:rPr')
    rFonts_fld = OxmlElement('w:rFonts')
    rFonts_fld.set(qn('w:ascii'), 'Arial')
    rFonts_fld.set(qn('w:hAnsi'), 'Arial')
    rPr_fld.append(rFonts_fld)
    sz_fld = OxmlElement('w:sz')
    sz_fld.set(qn('w:val'), '18')
    rPr_fld.append(sz_fld)
    color_fld = OxmlElement('w:color')
    color_fld.set(qn('w:val'), HEX_GREY_LIGHT)
    rPr_fld.append(color_fld)
    # Force non-bold on page number (prevent style inheritance in LibreOffice)
    b_off_fld = OxmlElement('w:b')
    b_off_fld.set(qn('w:val'), '0')
    rPr_fld.append(b_off_fld)

    r_begin = OxmlElement('w:r')
    r_begin.append(rPr_fld)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    r_begin.append(fldChar1)
    p_elem.append(r_begin)

    r_instr = OxmlElement('w:r')
    instrText = OxmlElement('w:instrText')
    instrText.text = ' PAGE '
    r_instr.append(instrText)
    p_elem.append(r_instr)

    r_end = OxmlElement('w:r')
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    r_end.append(fldChar2)
    p_elem.append(r_end)


# ---------------------------------------------------------------------------
# Paragraph helpers
# ---------------------------------------------------------------------------

def add_main_title(doc, text):
    """
    Main document title: Georgia 28pt bold BLUE_STEEL, centered.
    Used on the cover page only.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = BLUE_STEEL
    run.font.name = 'Georgia'
    return p


def add_section_heading(doc, text):
    """
    Section heading (Heading 1 equivalent): Georgia 16pt bold BLUE_STEEL.
    Built as a plain paragraph to avoid style interference on LibreOffice export.
    """
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = BLUE_STEEL
    run.font.name = 'Georgia'

    # Add spacing before heading via paragraph spacing
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '240')  # 12pt before
    spacing.set(qn('w:after'), '120')   # 6pt after
    pPr.append(spacing)

    return p


def add_sub_heading(doc, text):
    """
    Sub-heading (Heading 2 equivalent): Arial 12pt BLUE_DEEP.
    """
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.color.rgb = BLUE_DEEP
    run.font.name = 'Arial'

    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '120')
    spacing.set(qn('w:after'), '120')
    pPr.append(spacing)

    return p


def add_spacer(doc, count=1):
    """Add empty paragraphs as vertical spacers."""
    for _ in range(count):
        doc.add_paragraph()


def add_signature_block(doc, labels):
    """
    Fixed signature block (plain text, not a Content Control).
    Power Automate does not inject into this block.
    """
    add_spacer(doc)
    p_title = doc.add_paragraph()
    run = p_title.add_run(labels['title'])
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = GREY_DARK
    run.font.name = 'Arial'

    p_datum = doc.add_paragraph()
    p_datum.add_run(labels['datum']).font.size = Pt(10)

    p_sig = doc.add_paragraph()
    p_sig.add_run(labels['left_label']).font.size = Pt(10)
    p_sig.add_run("                    ").font.size = Pt(10)
    p_sig.add_run(labels['right_label']).font.size = Pt(10)

    p_line = doc.add_paragraph()
    p_line.add_run(labels['left_line']).font.size = Pt(10)
    p_line.add_run("       ").font.size = Pt(10)
    p_line.add_run(labels['right_line']).font.size = Pt(10)

    add_spacer(doc)


# ---------------------------------------------------------------------------
# Core build function
# ---------------------------------------------------------------------------

def build_template(output_path, strings):
    """
    Build a single .docx template with all 118 Content Controls and
    the 10 k Beratung visual identity.

    Structure:
      Cover page  : 6 SDTs  (COVER_TAGS)
      Profil      : 4 SDTs  (PROFIL_TAGS)
      Bilans 01-12: 9 SDTs each = 108 SDTs
      Total       : 118 SDTs
    """
    doc = Document()
    lang = strings.get("lang", "DE")

    # ---- Page setup ----
    section = doc.sections[0]
    section.page_width    = Cm(21)
    section.page_height   = Cm(29.7)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.0)

    # Base style: Georgia 11pt GREY_DARK
    doc.styles['Normal'].font.name = 'Georgia'
    doc.styles['Normal'].font.size = Pt(11)
    doc.styles['Normal'].font.color.rgb = GREY_DARK

    # ---- Header and footer ----
    set_header(section, strings["header_left"])
    set_footer(section, strings["footer_left"], lang=lang)

    # ====================================================================
    # COVER PAGE
    # ====================================================================

    add_main_title(doc, strings["doc_title_fixed"])
    add_spacer(doc, 2)

    cover_labels = strings["cover_labels"]
    table_cover = create_two_col_table(doc)

    for tag in COVER_TAGS:
        placeholder = strings["doc_titre_placeholder"] if tag == "doc_titre" else ""
        add_field_row(table_cover, cover_labels[tag], tag)

    add_spacer(doc, 2)

    conf_p = doc.add_paragraph(strings["confidential"])
    conf_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in conf_p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = GREY_LIGHT
        run.font.name = 'Arial'
        run.italic = True

    doc.add_page_break()

    # ====================================================================
    # SECTION 1: PROFIL
    # ====================================================================

    add_section_heading(doc, strings["section_profil_heading"])
    add_sub_heading(doc, strings["section_profil_sub"])
    add_spacer(doc)

    profil_labels = strings["profil_labels"]
    table_profil = create_two_col_table(doc)

    for tag in PROFIL_TAGS:
        add_field_row(table_profil, profil_labels[tag], tag)

    doc.add_page_break()

    # ====================================================================
    # SECTIONS BILANS 01 to 12
    # ====================================================================

    bilan_labels = strings["bilan_labels"]
    sig = strings["signature"]

    for nn in range(1, 13):
        prefix = f"bilan_{nn:02d}_"

        add_section_heading(doc, strings["bilan_heading"].format(nn=nn))
        add_sub_heading(doc, strings["bilan_sub"])
        add_spacer(doc)

        table_bilan = create_two_col_table(doc)
        for field in BILAN_FIELDS:
            tag = prefix + field
            add_field_row(table_bilan, bilan_labels[field], tag)

        add_signature_block(doc, sig)

        if nn < 12:
            doc.add_page_break()

    # ---- Save ----
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Saved: {output_path}")


# ---------------------------------------------------------------------------
# Verification helper
# ---------------------------------------------------------------------------

def verify_template(path, expected_count=118):
    """
    Count <w:tag w:val="..."> occurrences in document.xml and report.
    Each unique tag value should appear exactly once.
    """
    import zipfile
    import re

    with zipfile.ZipFile(path) as z:
        xml = z.read('word/document.xml').decode('utf-8')

    tags = re.findall(r'<w:tag\s+w:val="([^"]+)"', xml)

    duplicates = [t for t in set(tags) if tags.count(t) > 1]
    missing_cover  = [t for t in COVER_TAGS  if t not in tags]
    missing_profil = [t for t in PROFIL_TAGS if t not in tags]

    missing_bilans = []
    for nn in range(1, 13):
        for field in BILAN_FIELDS:
            tag = f"bilan_{nn:02d}_{field}"
            if tag not in tags:
                missing_bilans.append(tag)

    print(f"\nVerification: {path}")
    print(f"  Content Controls found : {len(tags)}")
    print(f"  Expected               : {expected_count}")
    print(f"  OK                     : {len(tags) == expected_count}")
    if duplicates:
        print(f"  WARN duplicates        : {duplicates}")
    if missing_cover:
        print(f"  MISSING cover tags     : {missing_cover}")
    if missing_profil:
        print(f"  MISSING profil tags    : {missing_profil}")
    if missing_bilans:
        print(f"  MISSING bilan tags     : {missing_bilans}")
    if not duplicates and not missing_cover and not missing_profil and not missing_bilans:
        print("  All 118 tag values present and unique.")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    base = os.path.dirname(os.path.abspath(__file__))

    de_path = os.path.join(base, "transfer_mappe_template_de.docx")
    en_path = os.path.join(base, "transfer_mappe_template_en.docx")

    print("Building DE template...")
    build_template(de_path, DE_STRINGS)

    print("Building EN template...")
    build_template(en_path, EN_STRINGS)

    print("\nRunning verification...")
    verify_template(de_path)
    verify_template(en_path)

    print("\nDone. Open the .docx files in Word to inspect Content Controls.")
    print("In Word: Developer tab > Design Mode to see all SDTs highlighted.")
